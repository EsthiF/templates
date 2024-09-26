import os
from docx import Document
from docx.shared import Pt


class WordTemplateProcessor:
    def __init__(self, template_path):
        """Initialize with the path to the template."""
        self.template_path = template_path  # Directly use the provided template path
        self.doc = None  # Placeholder for the document object

    def load_template(self):
        """Load the Word document template."""
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template not found: {self.template_path}")
        self.doc = Document(self.template_path)  # Load the Word document

    def replace_placeholders(self, placeholder_dict):
        """Replace placeholders in the document with the provided values."""
        if self.doc is None:
            raise ValueError("Document not loaded. Call 'load_template' before replacing placeholders.")
        
        # Replace placeholders in paragraphs
        for paragraph in self.doc.paragraphs:
            for placeholder, value in placeholder_dict.items():
                if placeholder in paragraph.text:
                    # Set right-to-left formatting
                    paragraph.paragraph_format.right_to_left = True

                    # Replace text in each run to preserve existing formatting
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            run.font.size=Pt(14)

        # Replace placeholders in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholder_dict.items():
                        if placeholder in cell.text:
                            # Set right-to-left formatting
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.right_to_left = True
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)

    def save_document(self, output_path):
        """Save the modified Word document to the specified output path."""
        if self.doc is None:
            raise ValueError("Document not loaded. Call 'load_template' before saving the document.")
        self.doc.save(output_path)
